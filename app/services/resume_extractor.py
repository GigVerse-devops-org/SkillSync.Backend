import logging
import pdfplumber
import docx
import io

logger = logging.getLogger(__name__)

class ResumeExtractionError(BaseException):
    pass

def extract_text_from_pdf(file_contents: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(file_contents)) as pdf:
            if not pdf.pages:
                raise ResumeExtractionError("PDF file contains no pages.")
            return "\n".join([page.extract_text() or "" for page in pdf.pages])
    except Exception as e:
        logger.error(f"Error extracting text from pdf: {str(e)}")
        raise ResumeExtractionError(f"Failed to extract text from pdf: {str(e)}")
        
def extract_text_from_docx(file_contents: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_contents))
        if not document.paragraphs:
            raise ResumeExtractionError("DOCX file contains no content.")
        return "\n".join([paragraph.text for paragraph in document.paragraphs])
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ResumeExtractionError(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_txt(file_contents: bytes) -> str:
    try:
        with io.BytesIO(file_contents) as f:
            text = f.read().decode("utf-8")
            if not text.strip():
                raise ResumeExtractionError("TXT file is empty.")
            return text
    except UnicodeDecodeError as e:
        logger.error("Failed to decode text as UTF-8")
        raise ResumeExtractionError("Invalid text encoding. Please ensure the file is UTF-8 encoded")
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        raise ResumeExtractionError(f"Failed to extract text from TXT: {str(e)}")

def extract_resume_text(file_contents: bytes, file_extension: str) -> str:
    if not file_contents:
        raise ResumeExtractionError("Not file contents provided.")
    
    file_extension = file_extension.lower()
    
    try:
        if file_extension == ".pdf":
            return extract_text_from_pdf(file_contents)
        elif file_extension in [".doc", ".docx"]:
            return extract_text_from_docx(file_contents)
        elif file_extension == ".txt":
            return extract_text_from_txt(file_contents)
        else:
            raise ResumeExtractionError(f"Unsupported file extension: {file_extension}")
    except ResumeExtractionError:
        raise
    except Exception as e:
        logger.error(f"Unexpected error during resume extraction: {str(e)}")
        ResumeExtractionError(f"Failed to process resume: {str(e)}")
    